import docx

def generate_cover_letter(company_name):
  cover_letter = f"""
  Dear Hiring Manager,

  I am excited to be applying for the Software Engineer Position at {company_name}. I graduated from the Computer Engineering program at University of Illinois at Urbana-Champaign in May 2020.  I have experience with front-end and back-end development in many different coding languages. I am comfortable working with AWS, GCP, and Azure Clouds. I have experience with using and building APIs and web applications. 

  At Optoro I worked in a microservices environment with a competitive tech stack of Ruby on Rails, Typescript, Node.js, Vue, and Angular. Optoro is a tech startup that focuses on the reverse logistics industry with clients like Target, Best Buy, and Gap. This experience has been incredibly valuable, and I have grown in my coding abilities.  I worked on a warehouse automation project enabling our warehouse to utilize robotics to improve efficiency and decrease costs. In addition, I lead a database consolidation project for our team to reduce our cloud footprint.  By consolidating many services databases into one database, we saved the company thousands of dollars. Following an initiative to improve metrics, I created a Grafana dashboard using Prometheus metrics for our various services. Furthermore, I worked on a project to update the versions of one of our tools including backend and frontend improvements. The Directed Sorting tool was used for efficiency improvements in the warehouse. This project was ultimately completed and released, and I learned many lessons in releasing software and supporting multiple versions. 

  Before working at Optoro, I built a website for a startup called Special Fitness using React.js hosted in AWS.  I have had one year of react experience developing this website.  Special Fitness helps people with disabilities workout by providing personal trainers.  I developed scheduling capabilities on this website between trainers and clients.  I served as the lead react developer on this project.  I used AWS services during development such as Amplify, DynamoDB, Cognito, Lambda, and IAM.

  Thank you for your consideration,
  Michael Gulson

  """

  
  return cover_letter


def create_word_doc(text, file_name):
  doc = docx.Document()
  doc.add_heading('Michael Gulson', 0)
  doc.add_paragraph('msgulson@gmail.com | (847)-323-8610 | 655 w Irving Park Rd Chicago, IL 60613')
  doc.add_paragraph(text)
  p = doc.add_paragraph(footer())

  hyperlink = add_hyperlink(p, 'https://github.com/michaelgulson', 'https://github.com/michaelgulson')
  hyperlink = add_hyperlink(p, 'https://www.linkedin.com/in/michael-gulson-578270111', 'https://www.linkedin.com/in/michael-gulson-578270111')
  hyperlink = add_hyperlink(p, 'https://gullman99.github.io/personal-website/', 'https://gullman99.github.io/personal-website/')

  doc.save(f"{file_name}.docx")

def footer():
  footer = """
  Github: https://github.com/michaelgulson
  LinkedIn Profile: https://www.linkedin.com/in/michael-gulson-578270111
  Personal Website: https://gullman99.github.io/personal-website/
  """
  return footer

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink



default_company = "Company"  # Default company name
company_name = input(f"Enter the company name (press enter to use the default {default_company}): ")
company_name = company_name.strip() if company_name else default_company

cover_letter_text = generate_cover_letter(company_name)
file_name = f"Cover_Letter(10-16-23){company_name.replace(' ', '_')}"
create_word_doc(cover_letter_text, file_name)
print(f"{file_name}.docx created successfully.")


